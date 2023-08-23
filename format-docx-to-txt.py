from docx import Document
import os

def refined_convert_docx_to_txt(docx_path, output_dir):
    """
    Convert a .docx file to a .txt file, capturing content from both paragraphs and tables.

    Parameters:
    - docx_path (str): Path to the .docx file.
    - output_dir (str): Directory to save the resulting .txt file.

    Returns:
    - str: Path to the generated .txt file.
    """
    
    # Create the output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Determine the name for the .txt file based on the .docx file name
    txt_filename = os.path.basename(docx_path).replace(".docx", ".txt")
    txt_path = os.path.join(output_dir, txt_filename)
    
    # Open the .docx file and write its content to the .txt file
    with open(txt_path, 'w', encoding='utf-8') as txt_file:
        doc = Document(docx_path)
        
        # Capture text from paragraphs
        for para in doc.paragraphs:
            txt_file.write(para.text + "\n")
        
        # Capture text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt_file.write(cell.text + "\n")
    
    return txt_path

# Example usage
source_dir = "../cv-docx"
output_dir = "../cv-txt"
docx_files = [os.path.join(source_dir, file) for file in os.listdir(source_dir) if file.endswith(".docx")]

for docx_file in docx_files:
    refined_convert_docx_to_txt(docx_file, output_dir)
